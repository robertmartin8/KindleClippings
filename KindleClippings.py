import re
import io
import os
import argparse
from fpdf import FPDF
import docx

def remove_chars(s, end_directory=""):
    """
    This is a utility function that removes special characters from the string, so that it can
    become a valid filename.
    :param s: input string
    :return: the input string, stripped of special characters
    """
    # Replace colons with a hyphen so "A: B" becomes "A - B"
    s = re.sub(" *: *", " - ", s)
    # Remove question marks or ampersands
    s = s.replace("?", "").replace("&", "and")
    # Replace ( ) with a hyphen so "this (text)" becomes "this - text"
    s = re.sub(r"\((.+?)\)", r"- \1", s)
    # Delete filename chars tht are not alphanumeric or ; , _ -
    s = re.sub(r"[^a-zA-Z\d\s;,_-]+", "", s)
    # Trim off anything that isn't a word at the start & end
    s = re.sub(r"^\W+|\W+$", "", s)

    max_length = 245 - len(end_directory)  # max file size limited to 255.
    s = s[:max_length]
    return s


def convert_to_format(path, file_name, format, encoding="utf-8"):
    """
    Will get text file and will convert to specified output

    :param path:
    :param file_name:
    :param format:
    :param encoding:
    :return: name of the file created
    """
    output_file_name = file_name[0:-4] + "." + format
    with io.open(path + file_name, "r", encoding=encoding, errors="ignore") as txt_file:

        paragraph = txt_file.read().split("\n")
        if format == "pdf":
            pdf_file = FPDF()
            pdf_file.add_page()
            # Construct absolute path for the font file
            script_dir = os.path.dirname(os.path.abspath(__file__))
            font_path = os.path.join(script_dir, 'media', 'NotoMono-Regular.ttf')
            pdf_file.add_font("mono", '', font_path, uni=True)
            pdf_file.set_font("mono", '', 11)

            for para in paragraph:
                # create muti-cell pdf object and add text to it
                # Using effective page width (A4 210mm - 2*10mm margin)
                pdf_file.multi_cell(190, 5, para, 0)
            pdf_file.output(path + output_file_name)

        elif format == "docx":
            docx_file = docx.Document()
            docx_file.add_heading(file_name[0:-4], 0)

            for para in paragraph:
                # add a paragraph and store the object in a variable
                docx_file.add_paragraph(para)
            docx_file.save(path + output_file_name)

    return output_file_name


def create_file_by_type(end_directory, format, encoding="utf-8"):
    """
    Will iterate over all text files and will convert and create file with specified format
    Currently Only pdf and docx are supported

    :param end_directory:
    :param format:
    :return: list of output filenames
    """
    output_files = []

    # get files in and directory
    files = [f for f in os.listdir(end_directory) if os.path.isfile(end_directory + f)]

    for file in files:
        if file[-3:] == "txt":
            output_files.append(convert_to_format(end_directory, file, format, encoding))

    return output_files


def parse_clippings(source_file, end_directory, encoding="utf-8", format="txt"):
    """
    Each clipping always consists of 5 lines:
    - title line
    - clipping info/metadata
    - a blank line
    - clipping text
    - a divider made up of equals signs
    Thus we can parse the clippings, and organise them by book.

    :param end_directory: the output directory where all of organised highlights will go
    :type end_directory: str
    :return: organises kindle highlights by book .
    """

    # Check that the source file (on the kindle) exists
    if not os.path.isfile(source_file):
        raise IOError("ERROR: cannot find " + source_file)

    # Create the output directory if it doesn't exist
    if not os.path.exists(end_directory):
        os.makedirs(end_directory)

    # This will keep track of the titles that we have already processed
    output_files = set()
    title = ""

    # Open clippings textfile and read data in lines
    with io.open(source_file, "r", encoding=encoding, errors="ignore") as f:
        # Individual highlights within clippings are separated by ==========
        for highlight in f.read().split("=========="):
            # For each highlight, we split it into the lines
            lines = highlight.split("\n")[1:]
            # Don't try to write if we have no body
            if len(lines) < 3 or lines[3] == "":
                continue
            # Set title and trim the hex character
            title = lines[0]
            if title[0] == "\ufeff":
                title = title[1:]

            # Remove characters and create path
            outfile_name = remove_chars(title, end_directory) + ".txt"
            path = end_directory + "/" + outfile_name

            # If we haven't seen title yet, set mode to write. Else, set to append.
            if outfile_name not in (list(output_files) + os.listdir(end_directory)):
                mode = "w"
                output_files.add(outfile_name)
                current_text = ""
            else:
                # If the title exists, read it as text so that we won't append duplicates
                mode = "a"
                with io.open(path, "r", encoding=encoding, errors="ignore") as textfile:
                    current_text = textfile.read()

            # Join all lines of the highlight text and strip leading/trailing whitespace
            clipping_text = "\n".join(lines[3:]).strip()

            with io.open(path, mode, encoding=encoding, errors="ignore") as outfile:
                # Write out the the clippings text if it's not already there
                if clipping_text not in current_text:
                    outfile.write(clipping_text + "\n\n...\n\n")

    # create additional file based on format
    if format in ["pdf","docx"]:
        formatted_out_files = create_file_by_type(end_directory, format, encoding)
        output_files.update(formatted_out_files)
    else:
        print("Invalid format mentioned. Only txt file will be created")
        args.format = "txt"

    print("\nExported titles:\n")
    for i in output_files:
        print(i)


if __name__ == "__main__":
    parser = argparse.ArgumentParser(
        description="Extract kindle clippings into a folder with nice text files"
    )
    parser.add_argument("-source", type=str, default="/Volumes/Kindle")
    parser.add_argument("-destination", type=str, default="./")
    parser.add_argument("-encoding", type=str, default="utf-8")
    parser.add_argument("-format", type=str, default="txt")
    args = parser.parse_args()

    if args.source[-4:] == ".txt":
        source_file = args.source
    elif args.source[-1] == "/":
        source_file = args.source + "documents/My Clippings.txt"
    else:
        source_file = args.source + "/documents/My Clippings.txt"

    if args.destination[-1] == "/":
        destination = args.destination + "KindleClippings/"
    else:
        destination = args.destination + "/KindleClippings/"

    parse_clippings(source_file, destination, args.encoding, args.format)
